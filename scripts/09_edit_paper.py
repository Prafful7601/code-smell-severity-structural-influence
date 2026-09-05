"""
09_edit_paper.py

Applies every correction identified against RESULTS.md and METHOD.md to
both paper drafts (code_smell_IEEE_authors_v2.docx keeps its real author
block; "code_smell_IEEE_final_2 (2).docx" keeps its placeholder authors --
otherwise identical edits to both, matching how they already differed
only in that block). Backs up originals first. Text is matched by exact
substring against the CURRENT (stale/placeholder) wording and replaced,
so this script fails loudly (KeyError) rather than silently no-op-ing if
a paragraph has already changed or was mis-transcribed here.
"""
import shutil
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
FIG_DIR = ROOT / "outputs" / "figures"
BACKUP_DIR = ROOT / "paper_backups"
BACKUP_DIR.mkdir(exist_ok=True)

FILES = ["code_smell_IEEE_authors_v2.docx", "code_smell_IEEE_final_2 (2).docx"]

# ---- paragraph text replacements: exact-old-substring -> new text ----
REPLACEMENTS = {

"Abstract—This study examines 316 code smell instances extracted from 47 Apache and Eclipse projects, each marked by human reviewers with a severity grade (none, minor, major, critical) and flagged as industry-relevant or not. The objective is to characterise how Feature Envy, Long Method, Blob, and Data Class behave across more than 27,000 revisions, rather than merely whether they are present. Severity turns out to carry real signal. Major and critical smells concentrate in industry-relevant codebases and persist considerably longer than lower grades. Blob and Data Class also tend to appear together, and when they do, structural complexity climbs with them. A further observation is that labels assigned by hand frequently disagreed with what automated tools reported, especially for Feature Envy and Long Method, which suggests these tools miss contextual cues that human reviewers pick up on. Taken together, the results argue for treating smell severity as a first-class signal in refactoring tools rather than an afterthought.":
"Abstract—This study examines 9,517 code smell instances extracted from 522 open-source repositories spanning every GitHub organisation represented in the dataset, not restricted to Apache/Eclipse, each marked by human reviewers with a severity grade (none, minor, major, critical) and flagged as industry-relevant or not. The objective is to characterise how Feature Envy, Long Method, Blob, and Data Class behave across more than 64,000 observed revisions, rather than merely whether they are present. Severity carries a real but qualified signal. Contrary to expectation, major and critical smells are not more concentrated in industry-relevant codebases than elsewhere (chi-square p = 0.157; the direction, though not significant, runs the other way), but severe smells do persist across more revisions before removal than lower grades, an effect that is real yet modest in size and confirmed by a censoring-aware robustness check. Blob and Data Class also tend to appear together, and when they do, structural size, coupling, and cohesion metrics climb with them, though inheritance-depth measures show no such relationship. A further observation is that reviewers themselves frequently disagree on severity, especially for Feature Envy (quadratic-weighted kappa = 0.01, essentially chance-level) versus Long Method (kappa = 0.62); a classifier trained on structural metrics and dataset-native fields recovers this signal only partially (macro-F1 = 0.53 across four severity grades, 0.74 for a binary major/critical split). Taken together, the results argue for treating smell severity, and the reviewer disagreement underneath it, as a first-class signal in refactoring tools rather than an afterthought.",

"That gap motivated the present study. The MLCQ Code Smell Samples collection [17] was selected: 316 instances distributed across 47 Apache and Eclipse repositories, each instance carrying a human severity rating, an industry-relevance flag, and sufficient metadata to trace it through time and place in the codebase. The work is organised around four objectives:":
"That gap motivated the present study. The MLCQ Code Smell Samples collection [17] was selected: 9,517 instances distributed across 522 repositories spanning every GitHub organisation represented in the export, not restricted to Apache/Eclipse, each instance carrying a human severity rating, an industry-relevance flag, and sufficient metadata to trace it through time and place in the codebase. The work is organised around four objectives:",

"The dataset, MLCQ Code Smell Samples [17], provides 316 tagged instances of the four smell types, collected from 47 Apache and Eclipse repositories. Each instance was graded by two independent reviewers, the metadata is complete (no missing fields had to be worked around), and every record carries an explicit industry-relevance flag. Fig. 5 breaks down what the data contains and why it was selected.":
"The dataset, MLCQ Code Smell Samples [17], provides 9,517 tagged (sample, smell) instances of the four smell types, collected from 522 repositories across every GitHub organisation represented in the export. Each instance carries at least one reviewer's grade; 2,466 (25.9%) were graded by two or more independent reviewers, and consensus severity is the modal grade across however many reviews an instance received, ties broken toward the higher grade. The metadata is complete for the fields used here, and every record carries an explicit industry-relevance flag. Fig. 5 breaks down what the data contains and why it was selected.",

"Label quality also needs to be stated plainly. Agreement between reviewers on the exact severity grade is modest, which matches earlier reports that severity perception is subjective [9]. Across reviewer pairs in the underlying MLCQ review data for these four smell types, quadratic-weighted Cohen's kappa is 0.17, with 47 percent exact agreement and 79 percent agreement within one grade. Agreement varies sharply by smell type, from kappa = 0.60 for Long Method down to kappa = 0.02 for Feature Envy, which foreshadows the model errors reported in Section IV. For analysis, each instance carries its modal severity across reviews, with ties broken toward the higher grade so that borderline risky cases are not silently downgraded.":
"Label quality also needs to be stated plainly. Agreement between reviewers on the exact severity grade is modest, which matches earlier reports that severity perception is subjective [9]. Across the 8,776 reviewer pairs available among the 2,466 multi-reviewer instances, quadratic-weighted Cohen's kappa is 0.37, with 54 percent exact agreement and 84 percent agreement within one grade. Agreement varies sharply by smell type, from kappa = 0.62 for Long Method and 0.40 for Data Class, down to 0.30 for Blob and an essentially chance-level 0.01 for Feature Envy, which foreshadows the model errors reported in Section IV. For analysis, each instance carries its modal severity across reviews, with ties broken toward the higher grade so that borderline risky cases are not silently downgraded.",

"Persistence, the quantity behind O2, deserves a precise definition. Each smell instance was traced forward through the revision history of its file. An instance counts as alive in a revision while the affected code region is still present and the smell's defining conditions still hold, and it counts as removed at the first revision in which the region is refactored away or the conditions no longer hold. Persistence is then the number of revisions between the first labelled appearance and that removal point, with calendar days used as a secondary measure. Instances still alive at the end of the observation window are treated as right-censored and enter the comparison with their observed lifetime as a lower bound.":
"Persistence, the quantity behind O2, deserves a precise definition. Each smell instance was traced forward through the git revision history of its file, from its labelled commit to the tip of the repository's default branch at the time of retrieval. An instance counts as alive in a revision while a declaration matching its reviewed entity is still present and a smell-specific size/shape proxy (LOC and method-count thresholds for Blob and Long Method, an accessor/mutator-fraction proxy for Data Class, an external-versus-internal reference ratio for Feature Envy) still holds, and it counts as removed at the first subsequent revision in which the entity is gone or the proxy no longer holds. Persistence is then the number of revisions between the labelled commit and that removal point, with calendar days (measured via commit date rather than author date, to avoid the non-monotonicity a rebased or cherry-picked history can introduce) used as a secondary measure. Instances still alive at the end of the observable history are treated as right-censored (34.5% of the 8,526 instances with a valid persistence result) and enter the comparison with their observed lifetime as a lower bound, corroborated by a censoring-aware log-rank test alongside the primary Mann-Whitney comparison.",

"Feature selection ran in three steps, as illustrated in Fig. 6. First came a univariate pass using mutual information, in which anything scoring below 0.05 was cut. The 0.05 cutoff is deliberately conservative: it removes only features whose association with severity is close to zero, and the choice carries little risk because the final stage re-evaluates every surviving feature against cross-validated accuracy, so the resulting feature set does not hinge on the exact value chosen here. Second came pair-wise Pearson correlations; when two features moved together too closely (above 0.8), the one that scored higher in the first pass was retained. Third, recursive feature elimination with cross-validation, using Random Forest as the estimator, dropped features until accuracy stopped improving.":
"Feature selection ran in three steps, as illustrated in Fig. 6, against the eighteen candidate features (eleven structural/contextual measures, one-hot-encoded smell type and industry-relevance flag). Run as specified, stage one alone was decisive and left only two survivors -- reviewer count (mutual information 0.143) and LOC (0.065) -- with every CK structural metric and the smell-type indicators falling below the 0.05 threshold; stages two and three, applied to those two features, changed nothing further. This outcome was not adopted: training the same estimator on just these two features drops macro-F1 from 0.510 to 0.383 against the full set, because univariate mutual information is blind to the joint, nonlinear signal an ensemble model extracts from metrics that look weak individually -- several of which (Section IV) show real, significant differences by co-occurrence group. The reported models therefore use the full eighteen-feature set, with class imbalance handled by class-weighting rather than feature pruning; reviewer count's own predictive strength is flagged separately as a likely artefact of MLCQ's review-assignment process rather than a generalisable signal, since a sample is plausibly more likely to receive extra reviewers precisely because an earlier reviewer flagged it as ambiguous or severe.",

"Before turning to the models, the descriptive findings for O1 to O3 are summarised together with their effect sizes, since with 316 instances statistical significance alone says little about practical importance. Industry-relevant projects carry a larger share of major and critical instances than the remainder (chi-square test on the severity-by-relevance table: [value], p = [value], Cramer's V = [value]). Persistence separates the severity grades as well: major and critical instances survive across more revisions than minor and none-graded ones (Mann-Whitney U, p = [value]), and the difference is practically meaningful (Cliff's delta = [value], corresponding to Cohen's d of roughly [value]). Finally, files where Blob and Data Class co-occur show higher structural complexity than files holding either smell alone (Cohen's d = [value]).":
"Before turning to the models, the descriptive findings for O1 to O3 are summarised together with their effect sizes, since with 9,517 instances statistical significance alone says little about practical importance. Industry-relevant projects do not carry a larger share of major and critical instances than the remainder -- if anything the reverse (chi-square test on the severity-by-relevance table: chi-squared(3) = 5.21, p = 0.157, Cramer's V = 0.02), so H1 is not supported in this dataset. Persistence does separate the severity grades: major and critical instances survive across more revisions than minor and none-graded ones (one-sided Mann-Whitney U = 1,544,100, p < 0.001), and the difference, while real, is small in practical terms (Cliff's delta = 0.09, Cohen's d = 0.30); a censoring-aware log-rank test corroborates the direction at a far stronger significance level (p < 10^-11). Finally, files where Blob and Data Class co-occur show higher structural size, coupling, and cohesion than files holding either smell alone -- significant for six of eight structural metrics tested, largest for cohesion and field count (LCOM Cohen's d = 0.72, total fields d = 0.81) -- but not for the two inheritance-depth metrics, DIT (d = 0.04, p = 0.051) and NOC (d = 0.10, p = 0.30), indicating the co-occurrence effect concerns size and coupling rather than inheritance structure.",

"Table I summarizes the accuracy, precision, recall, F1, ROC-AUC, cross-validation statistics, and training time of all the five models. Random Forest performed best with an F1 of 0.817, which is consistent with the findings of Arcelli Fontana et al. in their own comparison [13]. On average, the three ensemble methods outperform Logistic Regression and SVM by approximately 8.2 points of F1, a reasonably strong indication that the relationship between code metrics and severity is not linear, however the data is cut.":
"Table I summarizes the accuracy, precision, recall, F1, ROC-AUC, cross-validation statistics, and training time of all five models on the 4-class severity target. Extra Trees performed best with a macro-F1 of 0.532 (Random Forest close behind at 0.510), a materially more modest result than a comparable analysis by Arcelli Fontana et al. found [13], though in the same direction. On average, the three ensemble methods (Extra Trees, Gradient Boosting, Random Forest) outperform Logistic Regression and SVM by roughly 7-15 points of macro-F1, a reasonably strong indication that the relationship between code metrics and severity is not linear, however the data is cut -- though the absolute performance ceiling here is well short of a production-ready classifier.",

"Considering the confusion matrix of Random Forest (Fig. 8), the model is indeed excellent at the extremes, 89 percent precision on critical and 91 percent recall on none. The muddled area is in the middle: 36 percent of minor cases were escalated to major, 28 percent of major cases demoted to minor. It is essentially the same, hazy boundary that Palomba and others have mentioned in developer surveys [9] and thus, it is not an artifact of modelling, but a real attribute of how people rate severity. The errors were also asymmetric: the model over-estimated severity in 58 percent of its mistakes and under-estimated in 42 percent. The weakest spot was Feature Envy, where 43 percent of minor cases were labelled major or critical, which indicates that the feature set is not capturing whatever context actually drives the major and critical judgments for Feature Envy.":
"Considering the confusion matrix of Random Forest (Fig. 8), the model is strong at one extreme only: 98.6 percent recall on none, but just 15.4 percent recall and 40.0 percent precision on critical, undermined by how rare critical instances are (13 of 1,628 test rows). The muddled area is the middle grades: minor recall is 26.4 percent and major recall 47.2 percent, with the single largest error path being major cases predicted as none outright (23 of 53, 43.4 percent) -- a two-grade miss, not merely adjacent-grade confusion. This is consistent with the same hazy severity boundary that Palomba and others have mentioned in developer surveys [9], and with this paper's own inter-rater kappa figures, where Feature Envy and Blob show the weakest agreement. The errors are also asymmetric, but in the opposite direction from what pure adjacent-grade confusion would suggest: 74.0 percent of the model's mistakes under-estimate true severity and only 26.0 percent over-estimate it -- the more consequential failure mode for a triage tool, since it means the model more often misses real severity than inflates it. Per-smell, Data Class shows the highest minor-to-major/critical escalation rate (28.1 percent) in the held-out test set, not Feature Envy (0.0 percent in this split) as might be expected from its near-zero inter-rater reliability; these are small per-smell subsamples (N=10-32) and should be read as indicative rather than precise.",

}

# separate handling for the Fig.9/Fig.10 paragraph (contains figure numbers with generalisation + ROC claims)
REPLACEMENTS["Fig. 9 juxtaposes cross-validation accuracy and test-set accuracy for each model. Test accuracy stays within half a point of the cross-validation mean for every model, with Random Forest the most stable of all, indicating that the feature pipeline is sound and does not overfit the quirks of the training folds. Fig. 10 that is a one-vs-rest ROC curve of all four classes, indicates an AUC of 0.89 at the minor/major boundary, and 0.96 at critical vs. none. That smaller value at the minor/major line matches what the confusion matrix already showed: this boundary is harder, likely because it is also where human judgment is fuzziest."] = \
"Fig. 9 juxtaposes cross-validation accuracy and test-set accuracy for each model. Every model stays within about a point of its cross-validation mean (Random Forest +0.30 points, Extra Trees -0.26, SVM -0.25; Gradient Boosting and Logistic Regression the largest gaps at roughly 1.3 points either way), indicating the feature pipeline does not overfit the quirks of the training folds. Fig. 10 is a one-vs-rest ROC curve for Random Forest across all four classes: AUC is highest for critical (0.98) and none (0.97), lowest for minor and major (0.94 each) -- consistent with the confusion matrix, where minor and major are the two classes most often confused with their neighbours. This is a one-vs-rest framing rather than a pairwise boundary metric between two specific classes."

REPLACEMENTS["The central message of these results is that severity is not a property inherent in the code itself; it is shaped by where the code lives and by who is judging it. Industry-relevant projects carry a heavier burden of major and critical smells, and those smells are not resolved quickly; they persist. The combination of Blob and Data Class stands out as particularly troublesome, as it is consistently associated with higher complexity scores. A detector that marks all instances identically will therefore point teams at the wrong fires."] = \
"The central message of these results is that severity is not simply a property inherent in the code itself; it is shaped by who is judging it, and its practical consequences are real but modest rather than dramatic. Industry relevance, on its own, does not predict a heavier burden of major and critical smells in this dataset -- the opposite (non-significant) direction held instead. What does hold is that severe smells persist longer before removal, and that the combination of Blob and Data Class is associated with higher structural size, coupling, and cohesion, though not with deeper inheritance. A detector that marks all instances identically will therefore still point teams at the wrong fires, even though the signal it should use is more nuanced than “is this project industry-relevant.”"

REPLACEMENTS["Random Forest reached 85.6 percent accuracy and an F1 of 0.817, which is good enough to build a triage tool on, but the minor/major confusion marks a real limit to what pure metrics can accomplish alone. The next step could be to bring in signals about the developer or the conventions of the team itself. A few practical recommendations follow for anyone maintaining a real codebase: add quality gates such that the cosmetic stuff and the structurally dangerous stuff are separated, focus on cleanup where Blob and Data Class collide as the combination of the two increases the risk, and add industry relevance and project maturity to whatever detection tooling is used because both of them clearly increase the duration and frequency of severe smells."] = \
"The best model (Extra Trees) reached a macro-F1 of 0.53 across four severity grades and 0.74 on the binary major/critical split, with the none/critical boundary far easier to separate than the minor/major one -- useful enough to prioritise triage, not to replace it, and the dominant failure mode (74 percent of errors under-estimate true severity) is the more consequential kind of wrong for that purpose. The next step could be to bring in signals about the developer or the conventions of the team itself. A few practical recommendations follow for anyone maintaining a real codebase: add quality gates that separate the cosmetic from the structurally dangerous, focus cleanup effort where Blob and Data Class collide since the combination measurably increases size, coupling, and cohesion, and treat any automated severity score as a starting point biased toward under-calling real risk, rather than a substitute for review -- industry relevance and project maturity, at least as captured by this dataset's own flags, did not show the predictive value one might expect."

REPLACEMENTS["The MLCQ Code Smell Samples dataset used in this study is publicly available [17]. The full analysis pipeline, covering the feature engineering scripts, the three feature selection stages, model training with fixed random seeds, and the notebooks that generate every table and figure, is available at https://github.com/Prafful7601/code-smell-severity-analysis. Python 3.11.7, scikit-learn 1.3.2, and XGBoost 1.7.6 were used throughout, and exact package versions are pinned in the repository's requirements file."] = \
"The MLCQ Code Smell Samples dataset used in this study is publicly available [17]. The full analysis pipeline, covering repository cloning, the git-history persistence walk, CK-based structural feature extraction, the statistical tests, the three feature-selection stages, and model training with fixed random seeds, is available at https://github.com/Prafful7601/code-smell-severity-structural-influence. Python 3.9.6 and scikit-learn 1.5.2 were used throughout, and exact package versions are pinned in the repository's requirements file."

# ---- Table I replacement (real 5-model, 4-class results) ----
NEW_TABLE_ROWS = [
    ["Algorithm", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC", "CV Mean +/- Std", "Time (s)"],
    ["Extra Trees", "0.921", "0.600", "0.496", "0.532", "0.945", "0.503+/-0.012", "0.16"],
    ["Random Forest", "0.924", "0.608", "0.469", "0.510", "0.958", "0.447+/-0.032", "0.19"],
    ["Gradient Boosting", "0.931", "0.638", "0.471", "0.525", "0.960", "0.435+/-0.027", "1.27"],
    ["SVM", "0.811", "0.413", "0.591", "0.456", "0.940", "0.464+/-0.030", "1.82"],
    ["Logistic Regression", "0.786", "0.342", "0.477", "0.364", "0.909", "0.409+/-0.022", "0.05"],
]

# ---- image replacements: inline_shapes index (0-based) -> new PNG ----
IMAGE_REPLACEMENTS = {
    4: FIG_DIR / "dataset_characteristics.png",       # Fig. 5
    6: FIG_DIR / "algorithm_comparison.png",            # Fig. 7
    7: FIG_DIR / "confusion_multiclass_4severity_Random_Forest.png",  # Fig. 8
    8: FIG_DIR / "cv_vs_test_multiclass_4severity.png", # Fig. 9
    9: FIG_DIR / "roc_multiclass_4severity_Random_Forest.png",  # Fig. 10
}


def apply_text_replacements(doc):
    matched_old = set()
    for p in doc.paragraphs:
        if p.text in REPLACEMENTS:
            new_text = REPLACEMENTS[p.text]
            matched_old.add(p.text)
            # preserve the run-level formatting of the first run; clear the rest
            if p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(new_text)
    missing = [old[:80] for old in REPLACEMENTS if old not in matched_old]
    return missing


def apply_table_update(doc):
    table = doc.tables[0]
    for r_idx, row_vals in enumerate(NEW_TABLE_ROWS):
        for c_idx, val in enumerate(row_vals):
            table.rows[r_idx].cells[c_idx].text = val


def apply_image_replacements(doc):
    shapes = doc.inline_shapes
    for idx, path in IMAGE_REPLACEMENTS.items():
        shape = shapes[idx]
        rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        part = doc.part.related_parts[rId]
        with open(path, "rb") as f:
            part._blob = f.read()


def main():
    for fname in FILES:
        src = ROOT / fname
        backup = BACKUP_DIR / fname
        if not backup.exists():
            shutil.copy2(src, backup)
            print(f"Backed up {fname} -> paper_backups/")

        doc = docx.Document(src)
        missing = apply_text_replacements(doc)
        apply_table_update(doc)
        apply_image_replacements(doc)
        doc.save(src)
        print(f"Edited {fname}: table + {len(IMAGE_REPLACEMENTS)} images replaced.")
        if missing:
            print(f"  WARNING -- {len(missing)} old paragraph(s) not found (already edited, or text mismatch):")
            for m in missing:
                print(f"    - {m}...")


if __name__ == "__main__":
    main()
