"""
10_edit_authors.py

Author-block corrections requested after the initial push:
  - Neelam Rawat: add email + ORCID (previously had neither).
  - Mohd. Aatir: add email (previously had none).
  - New author: Dr. Prashant Agrawal (Associate Professor), inserted
    after Shweta Singh, matching the existing author-block formatting.

Applied to both docx files identically (author list is the same in both;
only the top-of-paper author BLOCK differs between the two files, which
this script does not touch -- "final_2 (2).docx" keeps its placeholder
author block untouched, per how the files have differed since before
this project started editing them).
"""
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parent.parent
AUTHOR_FILE = "code_smell_IEEE_authors_v2.docx"  # only this one has the real author block

RUN_SIZE = docx.shared.Pt(9)


def add_lines(paragraph, lines):
    for line in lines:
        r = paragraph.add_run("\n" + line)
        r.font.size = RUN_SIZE


def main():
    path = ROOT / AUTHOR_FILE
    doc = docx.Document(path)

    # locate by name rather than fixed index, in case the doc has shifted
    def find_author_para(name_substr):
        for p in doc.paragraphs:
            if p.style.name == "Author" and p.text.startswith(name_substr):
                return p
        raise ValueError(f"Author paragraph starting with {name_substr!r} not found")

    neelam = find_author_para("Neelam Rawat")
    assert "neelam.rawat@kiet.edu" not in neelam.text, "Neelam's email already present"
    add_lines(neelam, ["neelam.rawat@kiet.edu", "0000-0003-0759-6583"])
    print("Updated Neelam Rawat:", repr(neelam.text))

    aatir = find_author_para("Mohd. Aatir")
    assert "mohdaatir01@gmail.com" not in aatir.text, "Aatir's email already present"
    add_lines(aatir, ["mohdaatir01@gmail.com"])
    print("Updated Mohd. Aatir:", repr(aatir.text))

    shweta = find_author_para("Shweta Singh")

    # insert a new "Author"-style paragraph immediately after Shweta Singh
    new_p = doc.add_paragraph(style=shweta.style)
    shweta._p.addnext(new_p._p)

    lines = [
        "Dr. Prashant Agrawal",
        "Associate Professor",
        "Department of Computer Applications",
        "Krishna Institute of Engineering & Technology (KIET)",
        "Ghaziabad, Delhi-NCR, Uttar Pradesh, India",
        "prashant.agraw@gmail.com",
        "0000-0002-7890-024X",
    ]
    r = new_p.add_run(lines[0])
    r.font.size = RUN_SIZE
    add_lines(new_p, lines[1:])
    print("Inserted new author:", repr(new_p.text))

    doc.save(path)
    print(f"\nSaved {path}")


if __name__ == "__main__":
    main()
